#import functions
from docx import Document
from datetime import *
from middleware_dicts import *
import re as r
import os
import glob
import xlwt


def get_raw_tables_data (wordDoc):
    #this function get the data from the table without any adjust

    raw_data = []

    for table in wordDoc.tables:
        for row in table.rows:
            for cell in row.cells:
                raw_data.append(str(cell.text))

    raw_data = list(filter(None, raw_data))
    #print (raw_data)
    return (raw_data)

def get_tables_data (wordDoc):
    # get_tables_data extract data from the document's tables

    raw_data = []
    data = []

    for table in wordDoc.tables:
        for row in table.rows:
            for cell in row.cells:
                raw_data.append(str(cell.text))

    #print (raw_data)

    for i in range (len(raw_data)):

        if 'NOME DO PACIENTE:' in raw_data[i]:
            temp_data = raw_data[i] + raw_data[i+1]

        elif 'CONDIÇÃO DO INGRESSO:' in raw_data[i]:
            temp_data = raw_data[i] + raw_data[i+1]
            
        elif 'CONDIÇÃO DO EGRESSO:' in raw_data[i]:
            temp_data = raw_data[i] + raw_data[i+1]

        elif 'DN:' in raw_data[i]:
            temp_data = raw_data[i].replace('.', '/')

        elif 'DATA DO INGRESSO:' in raw_data[i]:
            temp_data = raw_data[i].replace('.', '/')

        elif 'DATA DA ALTA:' in raw_data[i]:
            temp_data = raw_data[i].replace('.', '/')
        
        elif 'DESLOCAMENTO:' in raw_data[i]:
            temp_data = 'DESLOCAMENTO: MANAUS'

        else:
            temp_data = raw_data[i]

        temp_data = temp_data.replace("\n", ' ').replace('  ',' ').replace('\t',' ')
        
        if ':' in temp_data:
            data.append(temp_data)

    #for i in data:
    #    print(i)
    #print(data)
    
    return (data)

def get_text (wordDoc):
    #get_text extract data from Document's texts

    fullText = []
    for para in wordDoc.paragraphs:
        fullText.append(para.text)
    new_fullText = list(filter(None, fullText))
    #return '\n'.join(new_fullText)
    #print (new_fullText)
    return (new_fullText)

def lowercase_text (fullText):
    #this function convert the text to lowercase and change latin caracters
    
    lc_fullText = []
    for i in fullText:
        lc_fullText.append(i.lower().replace('ç', 'c').replace('á', 'a').replace('é', 'e').replace('í', 'i').replace('ó', 'o').replace('ú', 'u').replace('ã', 'a').replace('ê', 'e'))
        
    #print(lc_fullText)
    return (lc_fullText)

def lowercase_table (tables_data):
    #this function convert the tables content to lowercase and change latin caracters

    lc_tablesdata = []
    for i in tables_data:
        lc_tablesdata.append(i.lower().replace('ç', 'c').replace('á', 'a').replace('é', 'e').replace('í', 'i').replace('ó', 'o').replace('ú', 'u').replace('ã', 'a').replace('ê', 'e'))
        
    #print(lc_tablesdata)
    return (lc_tablesdata)

def get_year (fullText):
    #get the year of the document

    new_text = lowercase_text(fullText)

    label = 'ANO: '

    for i in new_text:
        if 'manaus' in i[:10]:
            temp_year = i
            year = r.search(r"\d{4}", temp_year).group(0)
            return (label + year)

        #elif i[0].isdigit():
        #    temp_year = i
        #    year = r.search(r"\d{4}", temp_year).group(0)
        #    return (label + year)

    for i in range (len(new_text)):
        j = new_text[i].replace(' ', '')
        if 'manaus' in j[:10]:
            temp_year = j
            year = r.search(r"\d{4}", temp_year).group(0)
            return (label + year)

def get_month (fullText):
    #get the month of the document

    fullText = lowercase_text(fullText)

    month = 'MÊS: '
    for i in fullText:
        if 'manaus' in i:
            temp_month = i
            if 'janeiro' in temp_month:
                month = month + '01'
                return month
            elif 'fevereiro' in temp_month:
                month = month + '02'
                return month
            elif 'marco' in temp_month:
                month = month + '03'
                return month
            elif 'abril' in temp_month:
                month = month + '04'
                return month
            elif 'maio' in temp_month:
                month = month + '05'
                return month
            elif 'junho' in temp_month:
                month = month + '06'
                return month
            elif 'julho' in temp_month:
                month = month + '07'
                return month
            elif 'agosto' in temp_month:
                month = month + '08'
                return month
            elif 'setembro' in temp_month:
                month = month + '09'
                return month
            elif 'outubro' in temp_month:
                month = month + '10'
                return month
            elif 'novembro' in temp_month:
                month = month + '11'
                return month
            elif 'dezembro' in temp_month:
                month = month + '12'
                return month

def get_gender ():
    #this function insert a new collum in the sheet

    gender = 'SEXO: '
    return gender

def get_age (tables_data):
    #get age between birth and document date

    age_data = 'IDADE: '

    for i in tables_data:
        if 'DN:' in i:
            born_temp = str(r.findall(r':(.*)', i)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
            if (len(born_temp)) == 0:
                indices = [i for i, s in enumerate(tables_data) if 'DN:' in s]
                tables_data.insert(indices[0]+10, age_data)
                return (tables_data)
            else:    
                born = convert_year(born_temp)
                born = datetime.strptime(born, "%d/%m/%Y").date()
        if 'DATA DO INGRESSO:' in i:
            today_temp = str(r.findall(r':(.*)', i)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
            today = convert_year(today_temp)
            today = datetime.strptime(today, "%d/%m/%Y").date()

    age = today.year - born.year - ((today.month, today.day) < (born.month, born.day))

    
    age_data = age_data + str(age)
    
    #indices = [i for i, s in enumerate(tables_data) if 'DN:' in s]
    #tables_data.insert(indices[0]+10, age_data)

    tables_data.append(age_data)

    return (tables_data)

def get_entrydate (tables_data, text_data):
    #this function get the correct entrydate

    #print(tables_data)
    #print(lowercase_text(text_data))

    new_text = (lowercase_text(text_data))

    new_entry = 'DATA DO INGRESSO: '

    for i in range (len(tables_data)):
        if 'DATA DO INGRESSO:' in tables_data[i]:
            start_temp = str(r.findall(r':(.*)', tables_data[i])).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
            if (len(start_temp)) == 0:
                for j in new_text:
                    if j[0].isdigit():
                        #new_entry = new_entry + j[:10]
                        #new_entry = new_entry.replace('.','/')
                        entry = r.findall(r"\d{2}[./]\d{2}[./]\d{4}", j[:10])
                        if not entry:
                            entry = r.findall(r"\d{2}[./]\d{2}[./]\d{2}", j[:8])
                        new_entry = new_entry + entry[0]
                        new_entry = new_entry.replace('.','/')
                        tables_data[i] = new_entry
                        return (tables_data)

    return (tables_data)

def get_time (tables_data):
    #get time in hospital

    time_data = 'TEMPO NA CASAI (EM DIAS): '

    for i in tables_data:
        if 'DATA DO INGRESSO:' in i:
            start_temp = str(r.findall(r':(.*)', i)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
            init = convert_year(start_temp)
            init = datetime.strptime(init, "%d/%m/%Y").date()
        if 'DATA DA ALTA:' in i:
            end_temp = str(r.findall(r':(.*)', i)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
            if not end_temp:
                indices = [i for i, s in enumerate(tables_data) if 'DATA DA ALTA:' in s]
                tables_data.insert(indices[0]+1, time_data)
                return (tables_data)
            finish = convert_year(end_temp)
            finish = datetime.strptime(finish, "%d/%m/%Y").date()

    time = finish - init

    time_data = time_data + str(time.days)
    
    indices = [i for i, s in enumerate(tables_data) if 'DATA DA ALTA:' in s]
    tables_data.insert(indices[0]+1, time_data)

    return (tables_data)

def get_destination(tables_data):
    #get destination from-to

    for i in tables_data:
        if 'DESLOCAMENTO' in i:
            origin = str(r.findall(r':(.*)', i)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
        if 'PARA' in i:
            destiny = str(r.findall(r':(.*)', i)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
    destination_data = 'DESLOCAMENTO: '
    destination_data = destination_data + origin + '-' + destiny
    
    indices = [i for i, s in enumerate(tables_data) if 'PARA:' in s]
    tables_data.insert(indices[0]+1, destination_data)

    return (tables_data)

def convert_year(date):
    #this function convert year in format YY to YYYY

    #date = date.replace('.', '/')

    #date_temp = str(r.findall(r'/(.{2}$)', date)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
    date = date.replace(' ', '').replace(';','/').replace('.', '/')
    date_temp = date.split('/')[-1]
    date_temp = int(date_temp)
    if (len(str(date_temp))) <= 2:
        if date_temp < 10:
            complete = '200'
            date_temp = str(date_temp)
            date_new = complete + date_temp
        elif date_temp < 23:
            complete = '20'
            date_temp = str(date_temp)
            date_new = complete + date_temp
        else:
            complete = '19'
            date_temp = str(date_temp)
            date_new = complete + date_temp

        date = date[:-2] + date_new
        return (date)
    return (date)

def get_companion(tables_data):
    #This function return if the pacient have companion

    companion = 'ACOMPANHANTE: '
    for i in tables_data:
        if 'ACOMPANHANT' in i:
            lc_i = i.lower()
            companion_temp = str(r.findall(r':(.*)', lc_i)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
            if 'acompanhante' in companion_temp:
                #companion = companion + "N"
                companion = companion + "2"
                return companion
            elif len(companion_temp) >= 5:
                #companion = companion + "S"
                companion = companion + "1"
                return companion
            elif not companion_temp:
                #companion = companion + "N"
                companion = companion + "2"
                return companion        
            else:
                #companion = companion + "N"
                companion = companion + "2"
                return companion

def get_neglecteddiseases (tables_data, text_data):
    #this function search neglected diseases in document

    neglecteddiseases = [
        'malaria',
        'doenca de Chagas',
        'leishmaniose',
        'tuberculose',
        'tb',
        'dengue',
        'hanseniase',
        'esquistossomose',
        'oncocercose',
        'filariose',
        'tracoma',
        'helmintos',
        'nematoides de solo'
    ]

    new_text = lowercase_text(text_data)

    new_table = lowercase_table(tables_data)

    neglected = 'DOENÇA NEGLIGENCIADA: '

    for j in neglecteddiseases:
        for i in new_text:
            if j in i:
                #neglected = neglected + 'S'
                neglected = neglected + '1'
                return (neglected)
        for i in new_table:
            if j in i:
                #neglected = neglected + 'S'
                neglected = neglected + '1'
                return (neglected)
    
    #neglected = neglected + 'N'
    neglected = neglected + '2'
    return (neglected)

def get_neglecteddiseases_reason (tables_data, text_data):
    #this function search neglected diseases in document

    dict = {
        'malaria' : 'MALÁRIA',
        'doenca de chagas' : 'DOENÇA DE CHAGAS',
        'leishmaniose' : 'LEISHMANIOSE',
        'tuberculose' : 'TUBERCULOSE',
        'tb' : 'TUBERCULOSE',
        'dengue' : 'DENGUE',
        'hanseniase' : 'HANSENÍASE',
        'esquistossomose' : 'ESQUISTOSSOMOSE',
        'oncocercose' : 'ONCOCERCOSE',
        'filariose' : 'FILARIOSE',
        'tracoma' : 'TRACOMA',
        'helmintos' : 'HELMINTOS',
        'nematoides de solo' : 'NEMATÓIDES DE SOLO'
    }

    new_table = lowercase_table(tables_data)

    new_text = lowercase_text(text_data)

    #print(new_table)

    neglected_reason = 'MOTIVO NEGLIGENCIADA: '

    for i in new_table:
        for j in dict:
            if j in i:
                if neglected_reason.find(str(dict[j])) == -1 :
                    neglected_reason = neglected_reason + str(dict[j]) + '; '
    for i in new_text:
        for j in dict:
            if j in i:
                if neglected_reason.find(str(dict[j])) == -1 :
                    neglected_reason = neglected_reason + str(dict[j]) + '; '

    return neglected_reason

def get_conditionsensitive (dict, tables_data, text_data):
    #this function search disease sensitive to primary condition
    
    new_text = lowercase_text(text_data)

    new_table = lowercase_table(tables_data)

    conditionsensitive = 'DOENÇA SENSÍVEL À CONDIÇÃO PRIMÁRIA: '

    for i in new_table:
        for j in dict:
            if j in i:
                if ' ' in j:
                    conditionsensitive = conditionsensitive + '1'
                    return (conditionsensitive)
                else:
                    new_list = i.replace('.', '').replace(',', '').split(' ')
                    for i in new_list:
                        if i in j and len(i) == len(j):
                            conditionsensitive = conditionsensitive + '1'
                            return (conditionsensitive)
                #conditionsensitive = conditionsensitive + 'S'
                #conditionsensitive = conditionsensitive + '1'
                #return (conditionsensitive)
    for i in new_text:
        for j in dict:
            if j in i:
                if ' ' in j:
                    conditionsensitive = conditionsensitive + '1'
                    return (conditionsensitive)
                else:
                    new_list = i.replace('.', '').replace(',', '').split(' ')
                    for i in new_list:
                        if i in j and len(i) == len(j):
                            conditionsensitive = conditionsensitive + '1'
                            return (conditionsensitive)
                #conditionsensitive = conditionsensitive + 'S'
                #conditionsensitive = conditionsensitive + '1'
                #return (conditionsensitive)
    
    #conditionsensitive = conditionsensitive + 'N'
    conditionsensitive = conditionsensitive + '2'
    return (conditionsensitive)

def get_conditionsensitive_reason (dict, tables_data, text_data):

    new_table = lowercase_table(tables_data)

    new_text = lowercase_text(text_data)

    #print(new_table)

    conditionsensitive_reason = 'MOTIVO DOENÇA DE CONDIÇÃO PRIMÁRIA: '

    for i in new_table:
        for j in dict:
            if j in i:
                if ' ' in j:
                    if conditionsensitive_reason.find(str(dict[j])) == -1 :
                        conditionsensitive_reason = conditionsensitive_reason + str(dict[j]) + '; '
                else:
                    new_list = i.replace('.', '').replace(',', '').split(' ')
                    for i in new_list:
                        if i in j and len(i) == len(j):
                            if conditionsensitive_reason.find(str(dict[j])) == -1 :
                                conditionsensitive_reason = conditionsensitive_reason + str(dict[j]) + '; '
                #if conditionsensitive_reason.find(str(dict[j])) == -1 :
                #    conditionsensitive_reason = conditionsensitive_reason + str(dict[j]) + '; '
    for i in new_text:
        for j in dict:
            if j in i:
                if ' ' in j:
                    if conditionsensitive_reason.find(str(dict[j])) == -1 :
                        conditionsensitive_reason = conditionsensitive_reason + str(dict[j]) + '; '
                else:
                    new_list = i.replace('.', '').replace(',', '').split(' ')
                    for i in new_list:
                        if i in j and len(i) == len(j):
                            if conditionsensitive_reason.find(str(dict[j])) == -1 :
                                conditionsensitive_reason = conditionsensitive_reason + str(dict[j]) + '; '
                #if conditionsensitive_reason.find(str(dict[j])) == -1 :
                #    conditionsensitive_reason = conditionsensitive_reason + str(dict[j]) + '; '

    return conditionsensitive_reason

def get_giveup(text_data):
    #this function return if pacient give up or not

    giveup = 'DESISTÊNCIA: '

    lc_text = lowercase_text(text_data)

    for i in lc_text:
        if 'desist' in i:
            #giveup = giveup + "S"
            giveup = giveup + "1"
            return giveup

    #giveup = giveup + "N"
    giveup = giveup + "2"
    return giveup

def get_giveup_reason(text_data):
    #this function return the reason if the pacient had give up

    giveup_reason = 'MOTIVO DESIST: '

    lc_text = lowercase_text(text_data)

    for i in range (len(lc_text)):
        if 'desist' in lc_text[i]:
            giveup_reason = giveup_reason + str(text_data[i])
            return giveup_reason

    return giveup_reason

def get_internment(text_data):
    #this function return if pacient were internment or not

    internment = 'INTERNAÇÃO HOSPITALAR: '
    
    lc_text = lowercase_text(text_data)

    for i in lc_text:
        if 'interna' in i:
            #internment = internment + "S"
            internment = internment + "1"
            return internment

    #internment = internment + "N"
    internment = internment + "2"
    return internment

def get_referencedunit (dict, tables_data, text_data):
    #this function get all referenced units in document

    new_table = lowercase_table(tables_data)

    new_text = lowercase_text(text_data)

    #print(new_table)

    referencedunit = 'UNIDADE REFERENCIADA: '

    for i in new_table:
        for j in dict:
            if j in i:
                if ' ' in j:
                    if referencedunit.find(str(dict[j])) == -1 :
                        referencedunit = referencedunit + str(dict[j]) + '; '
                else:
                    new_list = i.replace('.', '').replace(',', '').split(' ')
                    for i in new_list:
                        if i in j and len(i) == len(j):
                            if referencedunit.find(str(dict[j])) == -1 :
                                referencedunit = referencedunit + str(dict[j]) + '; '
    for i in new_text:
        for j in dict:
            if j in i:
                if ' ' in j:
                    if referencedunit.find(str(dict[j])) == -1 :
                        referencedunit = referencedunit + str(dict[j]) + '; '
                else:
                    new_list = i.replace('.', '').replace(',', '').split(' ')
                    for i in new_list:
                        if i in j and len(i) == len(j):
                            if referencedunit.find(str(dict[j])) == -1 :
                                referencedunit = referencedunit + str(dict[j]) + '; '


    return referencedunit

def get_provdischarge (text_data):
    #this functions return if pacient had a provisional discharge

    provdischarge = 'ALTA PROVISÓRIA: '

    new_text = lowercase_text(text_data)

    for i in new_text:
        if 'paciente de alta provisoria para seu municipio de origem' in i:
            #provdischarge = provdischarge + "S"
            provdischarge = provdischarge + "1"
            return provdischarge
        if 'paciente segue de alta provisoria para seu municipio de origem' in i:
            #provdischarge = provdischarge + "S"
            provdischarge = provdischarge + "1"
            return provdischarge

    #provdischarge = provdischarge + "N"
    provdischarge = provdischarge + "2"
    return provdischarge

def get_problemsolved (tables_data):
    #this function return if the problem were solved

    problemsolved = 'PROBLEMA RESOLVIDO: '
    for i in tables_data:
        if 'ALTA PROVISÓRIA: N' in i:
            for i in tables_data:
                if 'DESISTÊNCIA: S' in i:
                    #problemsolved = problemsolved + 'N'
                    problemsolved = problemsolved + '2'
                if 'DESISTÊNCIA: N' in i:
                    #problemsolved = problemsolved + 'S'
                    problemsolved = problemsolved + '1'
        if 'ALTA PROVISÓRIA: S' in i:
            #problemsolved = problemsolved + 'N'
            problemsolved = problemsolved + '2'
        
    indices = [i for i, s in enumerate(tables_data) if 'ALTA PROVISÓRIA:' in s]
    tables_data.insert(indices[0]+1, problemsolved)

    return (tables_data)

def get_conditition (text_data):
    #this function return the pacient situation
    
    pacientcond = 'SITUAÇÃO DO PACIENTE: '

    #print(text_data)

    cont = []

    for i in text_data:
        if 'PENDENCIAS EM FILA DE ESPERA NO SISREG' in i:
            indice_1 = [i for i, s in enumerate(text_data) if 'PENDENCIAS EM FILA DE ESPERA NO SISREG' in s]
            cont.append(indice_1[0])
        if 'RONOGRAMA DE RETORNO CONSULTA/EXAME/CIRURGIA' in i:
            indice_2 = [i for i, s in enumerate(text_data) if 'CRONOGRAMA DE RETORNO CONSULTA/EXAME/CIRURGIA' in s]
            cont.append(indice_2[0])
        if 'TERAPIA MEDICAMENTOSA' in i:
            indice_3 = [i for i, s in enumerate(text_data) if 'TERAPIA MEDICAMENTOSA' in s]
            cont.append(indice_3[0])
        if 'CONSULTAS/EXAME/CIRURGIA' in i:
            indice_4 = [i for i, s in enumerate(text_data) if 'CONSULTAS/EXAME/CIRURGIA' in s]
            cont.append(indice_4[0])
        if 'REGISTRO DE INTERVENÇÕES' in i:
            indice_5 = [i for i, s in enumerate(text_data) if 'REGISTRO DE INTERVENÇÕES' in s]
            cont.append(indice_5[0])

    if (len(cont)) == 0:
        for i in text_data:
            if 'OBS:' in i:
                indice_extra = [i for i, s in enumerate(text_data) if 'OBS:' in s]
                cont.append(indice_extra[0]-1)

        if (len(cont)) == 0:
            for i in text_data:
                if 'RELATÓRIO DE CONTRA' in i:
                    indice_extra = [i for i, s in enumerate(text_data) if 'RELATÓRIO DE CONTRA' in s]
                    cont.append(indice_extra[0])

    indice = max(cont)

    pacientcond = pacientcond + str(text_data[indice+1:])

    return(pacientcond.replace('OBS:', 'OBS.').replace('\\t', '').replace("['",'').replace("', '", '').replace("']", '').replace('  ', ''))

def get_specialty (dict, tables_data, text_data):
    #this function get all the specialtys in document

    new_table = lowercase_table(tables_data)

    new_text = lowercase_text(text_data)

    #print(new_table)

    specialty = 'ESPECIALIDADES: '

    for i in new_table:
        for j in dict:
            if j in i:
                if specialty.find(str(dict[j])) == -1 :
                    specialty = specialty + str(dict[j]) + '; '
    for i in new_text:
        for j in dict:
            if j in i:
                if specialty.find(str(dict[j])) == -1 :
                    specialty = specialty + str(dict[j]) + '; '

    return specialty

def get_returndate (tables_data, raw_tablesdata, text_data):
    #this function get the date that the pacient must come back

    index_list = ['Data', 'Consulta', 'Médico', 'Local']

    index_lc = lowercase_text(index_list)

    tables_lc = lowercase_table(raw_tablesdata)

    index = []

    returndate = 'DATA DO RETORNO: '

    returnreason = 'MOTIVO RETORNO: '

    for i in text_data:
        if 'RONOGRAMA DE RETORNO CONSULTA' in i or 'RONOGRAMA DE RETORNO PARA CONSULTA' in i:
            for i in range(len(tables_lc)):
                #if tables_lc[i:i+len(index_lc)] == index_lc:
                    #index.append((i, i+len(index_list)))
                    #index.append((i+len(index_list)))
                #if index_lc[0] in tables_lc[i] and index_lc[1] in tables_lc[i+1] and index_lc[2] in tables_lc[i+2] and index_lc[3] in tables_lc[i+3]:
                if index_lc[0] in tables_lc[i] and index_lc[2] in tables_lc[i+2] and index_lc[3] in tables_lc[i+3]:
                    #index.append((i, i+len(index_list)))
                    index.append((i+len(index_list)))

            max_ind = max(index)

            if (len(raw_tablesdata)) > max_ind:
                date = r.findall(r"\d{2}[./]\d{2}[./]\d{4}", raw_tablesdata[max_ind])
                if not date:
                    date = r.findall(r"\d{2}[./]\d{2}[./]\d{2}", raw_tablesdata[max_ind])

                if date:
                    date = convert_year(date[0])

                if not date:
                    date.append(raw_tablesdata[max_ind])
                    date = date[0]

                date = str(date)

                returnreason = returnreason + str(raw_tablesdata[max_ind + 1])
            else:
                date = [' ']
                date = str(date)

            #date = str(date)

            returndate = returndate + date

            tables_data.append(returndate)

            tables_data.append(returnreason)

            return (tables_data)

    tables_data.append(returndate)

    tables_data.append(returnreason)

    return (tables_data)

def get_deltareturndate (tables_data):
    #get time between provisional discharge and return date

    time_data = 'TEMPO ALTA-RETORNO (EM DIAS): '

    for i in tables_data:
        if 'DATA DA ALTA:' in i:
            start_temp = r.findall(r"\d{2}[./]\d{2}[./]\d{4}", i)
            if not start_temp:
                start_temp = r.findall(r"\d{2}[./]\d{2}[./]\d{2}", i)
            if not start_temp:
                indices = [i for i, s in enumerate(tables_data) if 'DATA DO RETORNO:' in s]
                tables_data.insert(indices[0]+1, time_data)
                return (tables_data)
            init = convert_year(start_temp[0])
            init = datetime.strptime(init, "%d/%m/%Y").date()

        if 'DATA DO RETORNO:' in i:
            end_temp = r.findall(r"\d{2}[./]\d{2}[./]\d{4}", i)
            if not end_temp:
                indices = [i for i, s in enumerate(tables_data) if 'DATA DO RETORNO:' in s]
                tables_data.insert(indices[0]+1, time_data)
                return (tables_data)
            finish = convert_year(end_temp[0])
            finish = datetime.strptime(finish, "%d/%m/%Y").date()

    time = finish - init

    time_data = time_data + str(time.days)
    
    indices = [i for i, s in enumerate(tables_data) if 'DATA DO RETORNO:' in s]
    tables_data.insert(indices[0]+1, time_data)

    return (tables_data)

def get_path(file_path):
    #get path of the project and generate a hyperlink

    path = 'CAMINHO: file:///'
    path = path + str(file_path)
    path = path.replace("\\","/")
    return (path)

def get_ethnicity(tables_data, ethnicity_dict):
    #this function apply a fix in ethniticity column

    ethnicity = 'ETNIA: '

    lc_table = lowercase_table (tables_data)

    for i in lc_table:
        for j in ethnicity_dict:
            if j in i:
                if 'etnia:' in i:
                    ethnicity = ethnicity + str(ethnicity_dict[j])
                    #indice = [i for i, s in enumerate(tables_data) if 'ETNIA:' in s]
                    #tables_data.insert(indice[0]+1, ethnicity)

                    tables_data.append(ethnicity)
                    return (tables_data)

    return (tables_data)

def get_dsei(tables_data):
    #this function apply a fix in dsei column

    dsei_dict = {
        'alto rio negro' : 'ALTO RIO NEGRO',
        'arn' : 'ALTO RIO NEGRO',
        'alto rio solimoes' : 'ALTO RIO SOLIMÕES',
        'alto solimões' : 'ALTO RIO SOLIMÕES',
        'alto solimoes' : 'ALTO RIO SOLIMÕES',
        'manaus' : 'MANAUS',
        'medio purus' : 'MÉDIO RIO PURUS',
        'medio solimoes' : 'MÉDIO RIO SOLIMÕES',
        'parintins' : 'PARINTINS',
        'par' : 'PARINTINS',
        'vrj' : 'VALE DO JAVARI',
        'vale do javari' : 'VALE DO JAVARI',
        'arn/yan' : 'ALTO RIO NEGRO',
        'mp' : 'MÉDIO RIO PURUS',
        'ars' : 'ALTO RIO NEGRO',
        'alto rio solimões' : 'ALTO RIO SOLIMÕES',
        'mao' : 'MANAUS',
        'ms' : 'MÉDIO RIO SOLIMÕES',
        'medio rio purus' : 'MÉDIO RIO PURUS',
        'mrs' : 'MÉDIO RIO SOLIMÕES',
        'mrp' : 'MÉDIO RIO PURUS',
        'medio solimões' : 'MÉDIO RIO SOLIMÕES',
        'mrsa' : 'MÉDIO RIO SOLIMÕES E AFLUENTES',
        'vj' : 'VALE DO JAVARI',
        'yanomami' : 'YANOMAMI',
        'chico camilo' : '?',
        'alto rio negro.' : 'ALTO RIO NEGRO',
        'tonantins' : 'ALTO RIO SOLIMÕES',
        'alro solimões' : 'ALTO RIO SOLIMÕES',
        'alto soli mões' : 'ALTO RIO SOLIMÕES',
        'alvaraes' : 'MÉDIO RIO SOLIMÕES',
        'autazes' : 'MANAUS',
        'rmp' : 'MÉDIO RIO PURUS',
        'medio rio solimões' : 'MÉDIO RIO SOLIMÕES',
        'medio rio solimoes' : 'MÉDIO RIO SOLIMÕES',
        'm. rio solimoes' : 'MÉDIO RIO SOLIMÕES',
        'medio rio solimôes' : 'MÉDIO RIO SOLIMÕES',
        'medio solimoes.' : 'MÉDIO RIO SOLIMÕES',
        'vale do rio javari' : 'VALE DO JAVARI',
        'v. do javari' : 'VALE DO JAVARI'
    }

    dsei = 'DSEI DE ORIGEM: '

    lc_table = lowercase_table (tables_data)

    for i in lc_table:
        for j in dsei_dict:
            if j in i:
                if 'dsei de origem:' in i:
                    dsei = dsei + str(dsei_dict[j])
                    #indice = [i for i, s in enumerate(tables_data) if 'ETNIA:' in s]
                    #tables_data.insert(indice[0]+1, dsei)

                    tables_data.append(dsei)
                    return (tables_data)

    return (tables_data)

def get_to(tables_data):
    #this function apply a fix in para column

    to_dict = {
        'sao gabriel da cachoeira' : 'São Gabriel da Cachoeira',
        'sao gabreil da cachoeira' : 'São Gabriel da Cachoeira',
        'tabatinga' : 'Tabatinga',
        'autazes' : 'Autazes',
        'rio preto da eva' : 'Rio Preto da Eva',
        'nova olinda do norte' : 'Nova Olinda do Norte',
        'borba' : 'Borba',
        'urucara' : 'Urucará',
        'labrea' : 'Lábrea',
        'tefe' : 'Tefé',
        'maues' : 'Maués',
        'barreirinha' : 'Barreirinha',
        'atalaia do norte' : 'Atalaia do Norte',
        'santa isabel do rio negro' : 'Santa Isabel do Rio Negro',
        'barcelos' : 'Barcelos',
        'sao gabriel dacachoeira' : 'São Gabriel da Cachoeira',
        'sao gabriel da cachoeiro' : 'São Gabriel da Cachoeira',
        'franciane moreira' : '?',
        'benjamin constant' : 'Benjamin Constant',
        'sao paulo de olivenca' : 'São Paulo de Olivença',
        'santo antonio do ica' : 'Santo Antônio do Içá',
        'benjamim constant' : 'Benjamin Constant',
        'tonantins' : 'Tonantins',
        'amatura' : 'Amaturá',
        's.p. de olivenca' : 'São Paulo de Olivença',
        'manacapuru' : 'Manacapuru',
        'itacoatiara' : 'Itacoatiara',
        'manicore' : 'Manicoré',
        'silves' : 'Silves',
        'anama' : 'Anama',
        'beruri' : 'Beruri',
        'com. nsra da saude' : 'Manaus',
        'pb nsra da saude' : 'Manaus',
        'manaus' : 'Manaus',
        'careiro castanho' : 'Careiro Castanho',
        'polo base nsra da saude' : 'Manaus',
        'autazez' : 'Autazes',
        'nossa senhora da saude' : 'Manaus',
        'manaus / polo base' : 'Manaus',
        'careiro da varzea' : 'Careiro da Várzea',
        'polo base nossa senhora da saude' : 'Manaus',
        'taruma acu' : 'Manaus',
        'manaquiri' : 'Manaquiri',
        'jutai' : 'Jutaí',
        'tapaua' : 'Tapauá',
        'japura' : 'Japurá',
        'eirunepe' : 'Eirunepé',
        'jurua' : 'Juruá',
        'fonte boa' : 'Fonte Boa',
        'ipixuna' : 'Ipixuna',
        'maraa' : 'Maraã',
        'parintins' : 'Parintins',
        'santo antõnio do ica' : 'Santo Antônio do Içá',
        'n.o.n' : 'Nova Olinda do Norte',
        'barreitinha' : 'Barreirinha',
        'sgc' : 'São Gabriel da Cachoeira',
        'sao gabariel da cachoeira' : 'São Gabriel da Cachoeira',
        'boa vista do ramos' : 'Boa Vista do Ramos',
        'vale do javari' : '?',
        'manaus (paricatuba)' : 'Manaus',
        'itacotiara' : 'Itacoatiara',
        'carauari' : 'Carauari',
        'atalai ado norte' : 'Atalaia do Norte',
        'tocantins' : 'Tocantins',
        'santo antônio do ica' : 'Santo Antônio do Içá',
        'parinrins' : 'Parintins',
        'nhamunda' : 'Nhamundá',
        'sao gabriel da cahoeira' : 'São Gabriel da Cachoeira',
        'santa izabel do rio negro' : 'Santa Isabel do Rio Negro',
        'sao gabriel cachoeira' : 'São Gabriel da Cachoeira',
        'sao gabriel da cacheoira' : 'São Gabriel da Cachoeira',
        's ao gabriel da cachoeira' : 'São Gabriel da Cachoeira',
        'sao gabril da cachoeira' : 'São Gabriel da Cachoeira',
        'saõ gabriel da cachoeira' : 'São Gabriel da Cachoeira',
        'tonatins' : 'Tonantins',
        'santo antonio de ica' : 'Santo Antônio do Içá',
        'beijamin constant' : 'Benjamin Constant',
        'alvaraes' : 'Alvarães',
        'benjamin constante' : 'Benjamin Constant',
        'rio reto da eva' : 'Rio Preto da Eva',
        'rio preto d eva' : 'Rio Preto da Eva',
        'autazes / pantaleao' : 'Autazes',
        'manocore' : 'Manicoré',
        'novo olinda do norte' : 'Nova Olinda do Norte',
        'carreiro castanho' : 'Careiro Castanho',
        'novo airao' : 'Novo Airão',
        'canutama' : 'Canutama',
        'itamarati' : 'Itamarati',
        'coari' : 'Coari',
        'japura-tefe' : '?',
        'uarini' : 'Uarini',
        'eurunepe' : 'Eirunepé',
        'inhamunda' : 'Nhamundá',
        '' : ''
    }

    to = 'PARA: '

    lc_table = lowercase_table (tables_data)

    for i in lc_table:
        for j in to_dict:
            if j in i:
                if 'para:' in i:
                    to = to + str(to_dict[j])
                    #indice = [i for i, s in enumerate(tables_data) if 'ETNIA:' in s]
                    #tables_data.insert(indice[0]+1, to)

                    tables_data.append(to)
                    return (tables_data)

    return (tables_data)

def get_cid (tables_data, hd_dict):

    cid = 'CID: '

    lc_table = lowercase_table (tables_data)

    for i in lc_table:
        for j in hd_dict:
            if j in i:
                if cid.find(str(hd_dict[j])) == -1 :
                    cid = cid + str(hd_dict[j]) + '; '
    
                    #indice = [i for i, s in enumerate(tables_data) if 'HD:' in s]
                    #tables_data.insert(indice[0], cid)
                    tables_data.append(cid)

    if cid in 'CID: ':
        tables_data.append(cid)

    return (tables_data)

def get_transport(tables_data):
    #this function apply a fix in transport column

    transport_dict = {
        'fluvial' : 'fluvial',
        'ajato' : 'fluvial',
        'terrestre' : 'terrestre',
        'expresso' : 'fluvial',
        'barco com camarote/aereo ou a jato' : '?',
        'aereo' : 'aéreo',
        'aereo/fluvial.' : 'aéreo e fluvial',
        'aereo/barco(camarote)' : 'fluvial e aéreo',
        'aereo/lancha à jato' : 'fluvial e aéreo',
        'fluvial/aereo' : 'fluvial e aéreo',
        'fluvial /ajato' : 'fluvial',
        'barco' : 'fluvial',
        'rodoviario' : 'terrestre',
        'terrestre/fluvial' : 'terrestre e fluvial',
        'fluvial/terrestre' : 'fluvial e terrestre',
        'fluvial/aereo/terrestre' : 'fluvial, aéreo e terrestre',
        'terrestre/aereo/fluvial' : 'fluvial, aéreo e terrestre',
        'fluvial/aereo.' : 'fluvial e aéreo',
        'aereo, lancha a jato ou barco com camarote' : 'fluvial e aéreo',
        'f luvial/ terrestre' : 'terrestre e fluvial',
        'lancha ajato' : 'fluvial',
        'expesso' : 'fluvial',
        'via aereo' : 'aéreo',
        'expresso obs: paciente especial de colo com paralisia cerebral' : 'fluvial',
        'fluvial barco' : 'fluvial',
        'expresso obs.: paciente gravida' : 'fluvial',
        'expresso obs: paciente oncologico em tratamento' : 'fluvial',
        'fluvial/barco' : 'fluvial',
        'expresso obs: paciente gravida' : 'fluvial',
        'expresso obs.: paciente pos transplante de cornea' : 'fluvial',
        'expresso/ crianca de colo chorosa devido a patologia' : 'fluvial',
        'expresso / por conta do diagnostico de c.a' : 'fluvial',
        'expresso obs: paciente em uso de bolsa de colostomia' : 'fluvial',
        'barco de linha' : 'fluvial',
        'fluvial expresso' : 'fluvial',
        'expresso obs: crianca de colo cardiopata' : 'fluvial',
        'expresso obs: paciente idoso, pos cirurgico de amputacao de membro (cid n35)' : 'fluvial',
        'expresso paciente de colo com icc' : 'fluvial',
        'expresso obs: paciente cardiopata-especial' : 'fluvial',
        'expresso obs. com crianca de colo' : 'fluvial',
        'fluvial ou carona aerea' : '?',
        'expresso obs. paciente debilitado' : 'fluvial',
        'fluvial ou carona aereo' : 'fluvial',
        'expresso/ paciente em uso de medicacões injetaveis de 3 em 3 dias.' : 'fluvial',
        'expresso crianca realizou cirurgia' : 'fluvial',
        'expresso crianca realizou cirurgia.' : 'fluvial',
        'a criterio do dsei' : '?',
        'expresso/fluvial em leito ou aereo' : '?',
        '--' : '',
    }

    transport = 'MEIO DE TRANSPORTE: '

    lc_table = lowercase_table (tables_data)

    for i in lc_table:
        for j in transport_dict:
            if j in i:
                if 'meio de transporte:' in i:
                    transport = transport + str(transport_dict[j])
                    #indice = [i for i, s in enumerate(tables_data) if 'ETNIA:' in s]
                    #tables_data.insert(indice[0]+1, transport)

                    tables_data.append(transport)
                    return (tables_data)

    return (tables_data)

def get_scheduledreturn (tables_data):
    #this function return if return date were scheduled

    scheduledreturn = 'RETORNO AGENDADO: '

    for i in tables_data:
        if 'DATA DO RETORNO: ' in i:
            return_temp = str(r.findall(r':(.*)', i)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
            if (len(return_temp)) <= 4:
                #scheduledreturn = scheduledreturn + 'N'
                scheduledreturn = scheduledreturn + '2'
            else:
                #scheduledreturn = scheduledreturn + 'S'
                scheduledreturn = scheduledreturn + '1'            

    return (scheduledreturn)

def get_dataintext (argument):
    #this function get the names of ethnicity
    abs_path = os.path.abspath(os.curdir)

    files = glob.glob(abs_path + '/**/*.docx', recursive=True)

    ethnicity = argument.upper() + ':'

    f = open(argument.replace(' ', '_') + ".txt","w+")

    issues = 0
    invalid = 0

    argument = argument + ':'

    for file_path in range (len(files)):

        temp = files[file_path].split('\\')[-1]
    
        if '~$' in temp[:2]:
            issues = issues + 1
        if '$~' in temp[:2]:
            invalid = invalid + 1
        else:
            wordDoc = Document(files[file_path])
            #print(files[file_path])
            tables_data = get_tables_data(wordDoc)

            tables_data = lowercase_table(tables_data)
            
            for i in tables_data:
                if argument in i:
                        ethnicity_temp = str(r.findall(r':(.*)', i))
                        ethnicity_temp = ethnicity_temp.replace("[' ", "").replace("['", "").replace(" ']", "").replace("']", "")

                        if ethnicity.find(ethnicity_temp) == -1 :
                            ethnicity = ethnicity + ethnicity_temp + ';'

                            print(ethnicity_temp)
    
    
    ethnicity = ethnicity.replace(';', "' : ''\r")

    f.write ('%s' %ethnicity)

def get_dataintext_adult (argument):
    #this function get the names of ethnicity
    abs_path = os.path.abspath(os.curdir)

    files = glob.glob(abs_path + '/**/*.docx', recursive=True)

    ethnicity = argument.upper() + ':'

    f = open(argument.replace(' ', '_') + ".txt","w+")

    issues = 0
    invalid = 0

    argument = argument + ':'

    for file_path in range (len(files)):

        temp = files[file_path].split('\\')[-1]
    
        if '~$' in temp[:2]:
            issues = issues + 1
        if '$~' in temp[:2]:
            invalid = invalid + 1
        else:
            wordDoc = Document(files[file_path])
            #print(files[file_path])
            tables_data = get_tables_data(wordDoc)

            text_data = get_text(wordDoc)

            raw_tables_data = get_raw_tables_data (wordDoc)

            tables_data = get_entrydate(tables_data, text_data)

            tables_data = get_age(tables_data)

            tables_data.append(get_servicereceived (tables_data, raw_tables_data, text_data))

            tables_data = lowercase_table(tables_data)
            
            for j in tables_data:
                if 'idade:' in j:
                    age_temp = (r.findall(r':(.*)', j))
                    age_temp = age_temp[0].replace(' ', '')
                    if age_temp:
                        age_temp = int(age_temp)
                    if isinstance(age_temp, int):
                        if age_temp >= 18:
                            for i in tables_data:
                                if argument in i:
                                    ethnicity_temp = str(r.findall(r':(.*)', i))
                                    ethnicity_temp = ethnicity_temp.replace("[' ", "").replace("['", "").replace(" ']", "").replace("']", "")

                                    if ethnicity.find(ethnicity_temp) == -1 :
                                        ethnicity = ethnicity + ethnicity_temp + ';'

                                        print(ethnicity_temp)
    
    
    ethnicity = ethnicity.replace(';', "' : ''\r")

    f.write ('%s' %ethnicity)

def get_servicereceived (tables_data, raw_tables_data, text_data, dict) :
    #get all the services received during the internment

    index_lc = ['data', 'consulta', 'medico', 'local']

    index_lc_ = ['data', 'exames realizados', 'local']

    endindex_lc = ['data', 'medicamento', 'tratamento']

    tables_lc = lowercase_table(raw_tables_data)

    end_index = [len(tables_lc)]

    index = []

    examlist = 'ATENDIMENTO RECEBIDO: '

    examlist_new = 'ATENDIMENTOS RECEBIDOS: '

    for i in tables_data:
        if 'DATA DA ALTA:' in i:
            finish = ''
            end_temp = str(r.findall(r':(.*)', i)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
            if not end_temp:
                break
            finish = convert_year(end_temp)
            finish = datetime.strptime(finish, "%d/%m/%Y").date()

            break

    for j in text_data:
        if 'REGISTRO DE INTERVEN' in j:
            for i in range(len(tables_lc)):
                if index_lc_[0] in tables_lc[i] and index_lc_[1] in tables_lc[i+1] and index_lc_[2] in tables_lc[i+2]:
                    index.append((i+len(index_lc_)))

        if 'CONSULTAS/EXAME/CIRURGIA' in j:
            for i in range(len(tables_lc)):
                if index_lc[0] in tables_lc[i] and index_lc[2] in tables_lc[i+2] and index_lc[3] in tables_lc[i+3]:
                    index.append((i+len(index_lc)))

    for i in range(len(tables_lc)):
        if endindex_lc[0] in tables_lc[i] and endindex_lc[1] in tables_lc[i+1] and endindex_lc[2] in tables_lc[i+2]:
            end_index.append(i)
        if 'protocolo' in tables_lc[i]:
            end_index.append(i)

    if not index:
        index = [len(tables_lc)]
        min_ind = min(index)
    else:
        min_ind = min(index)

    max_ind = min(end_index)

    examslist_temp = tables_lc[min_ind:max_ind]

    for i in range (len(examslist_temp)):
        date_temp = examslist_temp[i]
        if date_temp[:2].isdigit():
            date_new = r.findall(r"\d{2}[./]\d{2}[./]\d{4}", date_temp)
            if not date_new:
                date_new = r.findall(r"\d{2}[./]\d{2}[./]\d{2}", date_temp)
            if not date_new:
                continue
            exam_date = convert_year(date_new[0])
            exam_date = datetime.strptime(exam_date, "%d/%m/%Y").date()

            if not finish:
                break

            if finish > exam_date:
                if not 'ista' in examslist_temp[i+1]:
                    #print(examslist_temp[i+1])

                    if examlist.find(examslist_temp[i+1]) == -1 :
                        examlist = examlist + str(examslist_temp[i+1]).replace('\t', '') + '; '
                for j in dict:
                    if j in examslist_temp[i+1]:
                        if examlist_new.find(str(dict[j])) == -1 :
                            examlist_new = examlist_new + str(dict[j]) + '; '
                        if examlist.find(str(dict[j])) == -1 :
                            examlist = examlist + str(examslist_temp[i+1]) + '; '
    
    tables_data.append(examlist_new)

    tables_data.append(examlist)

    return (tables_data)

def get_examsperformed (argument):
    #this function get the exams performed during the interment

    abs_path = os.path.abspath(os.curdir)

    files = glob.glob(abs_path + '/**/*.docx', recursive=True)

    examlist = argument.upper() + ': '

    f = open(argument.replace(' ', '_') + ".txt","w+")

    issues = 0
    invalid = 0

    for file_path in range (len(files)):

        temp = files[file_path].split('\\')[-1]
    
        if '~$' in temp[:2]:
            issues = issues + 1
        if '$~' in temp[:2]:
            invalid = invalid + 1
        else:
            wordDoc = Document(files[file_path])
            #print(files[file_path])

            tables_data = get_tables_data(wordDoc)

            raw_tables_data = get_raw_tables_data (wordDoc)

            text_data = get_text(wordDoc)

            #column_number: 1 - consulta, 2 - medico, 3 - local

            index_lc = ['data', 'consulta', 'medico', 'local']

            index_lc_ = ['data', 'exames realizados', 'local']

            endindex_lc = ['data', 'medicamento', 'tratamento']

            tables_lc = lowercase_table(raw_tables_data)

            #text_lc = lowercase_text(text_data)

            index = []

            end_index = [len(tables_lc)]

            for i in tables_data:
                if 'DATA DA ALTA:' in i:
                    end_temp = str(r.findall(r':(.*)', i)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
                    if not end_temp:
                        break
                    finish = convert_year(end_temp)
                    finish = datetime.strptime(finish, "%d/%m/%Y").date()

                    break

            for j in text_data:
                if 'REGISTRO DE INTERVEN' in j:
                    for i in range(len(tables_lc)):
                        if index_lc_[0] in tables_lc[i] and index_lc_[1] in tables_lc[i+1] and index_lc_[2] in tables_lc[i+2]:
                            index.append((i+len(index_lc_)))

                if 'CONSULTAS/EXAME/CIRURGIA' in j:
                    for i in range(len(tables_lc)):
                        if index_lc[0] in tables_lc[i] and index_lc[2] in tables_lc[i+2] and index_lc[3] in tables_lc[i+3]:
                            index.append((i+len(index_lc)))

            for i in range(len(tables_lc)):
                if endindex_lc[0] in tables_lc[i] and endindex_lc[1] in tables_lc[i+1] and endindex_lc[2] in tables_lc[i+2]:
                    end_index.append(i)
                if 'protocolo' in tables_lc[i]:
                    end_index.append(i)

            if not index:
                index = [len(tables_lc)]
                min_ind = min(index)
            else:
                min_ind = min(index)

            max_ind = min(end_index)

            examslist_temp = tables_lc[min_ind:max_ind]

            for i in range (len(examslist_temp)):
                date_temp = examslist_temp[i]
                if date_temp[:2].isdigit():
                    date_new = r.findall(r"\d{2}[./]\d{2}[./]\d{4}", date_temp)
                    if not date_new:
                        date_new = r.findall(r"\d{2}[./]\d{2}[./]\d{2}", date_temp)
                    if not date_new:
                        continue
                    exam_date = convert_year(date_new[0])
                    exam_date = datetime.strptime(exam_date, "%d/%m/%Y").date()

                    if not finish:
                        break

                    if finish > exam_date:
                        if not 'ista' in examslist_temp:
                            #print(examslist_temp)

                            for i in range (len(examslist_temp)):
                                if not '18' in examslist_temp[i] and not '19' in examslist_temp[i] and not '20' in examslist_temp[i] and not '21' in examslist_temp[i] and not '17' in examslist_temp[i]:
                                    if examlist.find(str(examslist_temp[i])) == -1 :
                                        print(examslist_temp[i])
                                        examlist = examlist + str(examslist_temp[i]) + ';'

    examlist = examlist.replace(';', "' : ''\r")

    f.write ('%s' %examlist)

def get_outputlog (list, issues, invalid, valid):
    #create a log the projet's root with a short resume of document's status

    f= open("output_log.txt","w+")

    f.write ('**********************************************************\r')
    f.write ('\nvalid files: (%d file(s))\n\r' %valid)

    f.write ('**********************************************************\r')
    f.write ('\nlist of corrupted files: (%d file(s))\n\r' %issues)

    for i in range (len(list)):
        if '~$' in list[i]:
            f.write ("%s\r" % list[i])

    f.write ('\n**********************************************************\r')
    f.write ('\nlist of invalid files: (%d file(s))\n\r' %invalid)

    for i in range (len(list)):
        if '$~' in list[i]:
            f.write ("%s\r" % list[i])

    f.close()

def organizer (tables_data):
    #organize the colluns

    #print(tables_data)

    new_table = ['s'] * 39
    
    for i in tables_data:
        if 'ANO:' in i:
            new_table[0] = i
        elif 'MÊS:' in i:
            new_table[1] = i
        elif 'DSEI DE ORIGEM:' in i:
            new_table[2] = i
        elif 'SEXO:' in i:
            new_table[3] = i
        elif 'NOME DO PACIENTE:' in i:
            new_table[4] = i
        elif 'DN:' in i:
            new_table[5] = i
        elif 'IDADE:' in i:
            new_table[6] = i
        if 'COMUNIDADE:' in i:
            new_table[7] = i
        elif 'ETNIA:' in i:
            new_table[8] = i
        elif 'DATA DO INGRESSO:' in i:
            new_table[9] = i
        elif 'DATA DA ALTA:' in i:
            new_table[10] = i
        elif 'TEMPO NA CASAI (EM DIAS):' in i:
            new_table[11] = i
        elif 'RETORNO AGENDADO:' in i:
            new_table[12] = i
        elif 'DATA DO RETORNO:' in i:
            new_table[13] = i
        elif 'TEMPO ALTA-RETORNO (EM DIAS):' in i:
            new_table[14] = i
        elif 'MOTIVO RETORNO:' in i:
            new_table[15] = i
        elif 'CID:' in i:
            new_table[16] = i
        elif 'HD:' in i:
            new_table[17] = i
        elif 'ESPECIALIDADES:' in i:
            new_table[18] = i
        elif 'CONDIÇÃO DO INGRESSO:' in i:
            new_table[19] = i
        elif 'CONDIÇÃO DO EGRESSO:' in i:
            new_table[20] = i
        elif 'INTERNAÇÃO HOSPITALAR:' in i:
            new_table[21] = i
        elif 'ATENDIMENTOS RECEBIDOS:' in i:
            new_table[22] = i
        elif 'ATENDIMENTO RECEBIDO:' in i:
            new_table[23] = i
        elif 'UNIDADE REFERENCIADA:' in i:
            new_table[24] = i
        elif 'DESLOCAMENTO:' in i:
            new_table[25] = i
        elif 'PARA:' in i:
            new_table[26] = i
        elif 'MEIO DE TRANSPORTE:' in i:
            new_table[27] = i
        elif 'ACOMPANHANTE:' in i:
            new_table[28] = i
        elif 'ALTA PROVISÓRIA:' in i:
            new_table[29] = i
        elif 'DOENÇA NEGLIGENCIADA:' in i:
            new_table[30] = i
        elif 'MOTIVO NEGLIGENCIADA:' in i:
            new_table[31] = i
        elif 'DOENÇA SENSÍVEL' in i:
            new_table[32] = i
        elif 'MOTIVO DOENÇA DE CONDI' in i:
            new_table[33] = i
        elif 'SITUAÇÃO DO PACIENTE:' in i:
            new_table[34] = i
        elif 'PROBLEMA RESOLVIDO:' in i:
            new_table[35] = i
        elif 'DESISTÊNCIA:' in i:
            new_table[36] = i
        if 'MOTIVO DESIST:' in i:
            new_table[37] = i
        elif 'CAMINHO:' in i:
            new_table[38] = i
    return new_table

def get_data (wordDoc, ethnicity_dict, spec_dict, sensitive_dict, servicereceived_dict, hd_dict, hospital_dict , file_path):
    #generate tables_data
    
    tables_data = get_tables_data(wordDoc)

    raw_tables_data = get_raw_tables_data (wordDoc)

    text_data = get_text(wordDoc)

    tables_data.insert(0, get_year(text_data))

    tables_data.insert(1, get_month(text_data))

    tables_data.insert(2, get_gender())

    tables_data = get_entrydate(tables_data, text_data)

    tables_data = get_ethnicity(tables_data, ethnicity_dict)

    tables_data = get_dsei(tables_data)

    tables_data = get_to(tables_data)

    tables_data = get_cid(tables_data, hd_dict)

    tables_data = get_transport(tables_data)

    tables_data = get_age(tables_data)

    tables_data = get_time(tables_data)

    tables_data.append(get_specialty(spec_dict, raw_tables_data, text_data))

    tables_data.append(get_companion(tables_data))

    tables_data.append(get_provdischarge(text_data))

    tables_data.append(get_giveup(text_data))

    tables_data.append(get_giveup_reason(text_data))

    tables_data = get_problemsolved(tables_data)

    tables_data.append(get_neglecteddiseases(tables_data, text_data))

    tables_data.append(get_neglecteddiseases_reason(tables_data, text_data))

    tables_data.append(get_conditionsensitive(sensitive_dict, tables_data, text_data))

    tables_data.append(get_conditionsensitive_reason(sensitive_dict, tables_data, text_data))

    tables_data.append(get_conditition(text_data))

    tables_data.append(get_internment(text_data))

    tables_data = get_servicereceived (tables_data, raw_tables_data, text_data, servicereceived_dict)

    tables_data.append(get_referencedunit(hospital_dict, tables_data, text_data))

    tables_data = get_returndate(tables_data, raw_tables_data, text_data)

    tables_data.append(get_scheduledreturn(tables_data))

    tables_data = get_deltareturndate(tables_data)

    tables_data.append(get_path(file_path))

    new_table = organizer(tables_data)

    return new_table

def run_automation():
    #run the automation tool
    abs_path = os.path.abspath(os.curdir)

    files = glob.glob(abs_path + '/**/*.docx', recursive=True)

    bad_files = glob.glob(abs_path + '/**/*.doc', recursive=True)

    issues = 0
    valid = 0
    invalid = 0

    book = xlwt.Workbook(encoding="utf-8")

    sheet1 = book.add_sheet("Sheet 1")

    style = xlwt.XFStyle()
    font = xlwt.Font()
    font.bold = True
    style.font = font

    specs = []
    infos = []

    #all_tables_data = [[0]*(len(files))]*26
    for file_path in range (len(files)):

        temp = files[file_path].split('\\')[-1]
    
        if '~$' in temp[:2]:
            issues = issues + 1
        if '$~' in temp[:2]:
            invalid = invalid + 1

        else:
            wordDoc = Document(files[file_path])
            #print(files[file_path])
            tables_data = get_data(wordDoc, ethnicity_dict, specialist_dict, conditionsensitive_dict, servicereceived_dict, hd_dict, hospital_dict, files[file_path])

            for j in tables_data:
                specs_temp = str(r.findall(r'(.*):', j)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
                infos_temp = str(r.findall(r':(.*)', j)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
                specs.append(specs_temp)
                infos.append(infos_temp)

            for j in range (len(tables_data)):
                if valid == 0:
                    sheet1.write(0, j, specs[j], style=style)
                    sheet1.write((valid+1), j, infos[j])
                else:
                    sheet1.write((valid+1), j, infos[j])

            del specs[:]
            del infos[:]

            valid = valid + 1
            
            os.system('cls')

            print('%d of ' %valid, (len(files)))
            #for columns in range (len(tables_data)):
            #    all_tables_data[file_path][columns] = tables_data[columns]

    for file_path in range (len(bad_files)):
        temp = bad_files[file_path].split('\\')[-1]
        if '~$' in temp[:2]:
            issues = issues + 1
        if '$~' in temp[:2]:
            invalid = invalid + 1

    sheet1.col(0).width = 1400
    sheet1.col(1).width = 1400
    sheet1.col(2).width = 5000
    sheet1.col(3).width = 1400
    sheet1.col(4).width = 7000
    sheet1.col(5).width = 2600
    sheet1.col(6).width = 2000
    sheet1.col(7).width = 5000
    sheet1.col(8).width = 4000
    sheet1.col(9).width = 2600
    sheet1.col(10).width = 2600
    sheet1.col(11).width = 4000
    sheet1.col(12).width = 4000

    now = datetime.now()
    dt_string = now.strftime("%d%m%Y_%H%M%S")

    book.save("output_" + dt_string + ".xls")
    
    list = []
    list = files
    list.extend(bad_files)

    get_outputlog(list, issues, invalid, valid)

    #book.save("test.xls")

    print ('\n**********************************************************')
    print ('There is %d corrupted files!' %issues)
    print ('There is %d invalid files!' %invalid)
    print ('There is a total of %d valid files!' %valid)
    print ('**********************************************************')

    return (0)

def create_sheet (data):
    #this function create a sheet
    
    book = xlwt.Workbook(encoding="utf-8")

    sheet1 = book.add_sheet("Sheet 1")

    specs = []
    infos = []

    for i in data:
        specs_temp = str(r.findall(r'(.*):', i)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
        infos_temp = str(r.findall(r':(.*)', i)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
        specs.append(specs_temp)
        infos.append(infos_temp)

    for i in range (len(data)):
        sheet1.write(0, i, specs[i])
        sheet1.write(1, i, infos[i])

    book.save("test.xls")
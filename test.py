from docx import Document
from datetime import *
import re
import os

print(os.path.abspath(os.getcwd()))

wordDoc = 'CONTRA REF 2019\MAO\OUTUBRO\ARGEL PALHETA DE MENEZES.docx'

def get_tables_data (wordDoc):
    raw_data = []
    data = []

    for table in wordDoc.tables:
        for row in table.rows:
            for cell in row.cells:
                raw_data.append(str(cell.text))

    #print (raw_data)

    for i in range (len(raw_data)):
        if 'NOME DO PACIENTE:' in raw_data[i]:
            temp_data = raw_data[i] + raw_data[i+1]
        else:
            temp_data = raw_data[i]

        if ':' in temp_data:
            data.append(temp_data)

    #for i in data:
    #    print(i)
    
    return (data)

def getText(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    #return '\n'.join(fullText)
    return (fullText)

def get_destination(tables_data):
    for i in tables_data:
        if 'DESLOCAMENTO' in i:
            origin = str(r.findall(r':(.*)', i)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
        if 'PARA' in i:
            destiny = str(r.findall(r':(.*)', i)).replace("[' ", '').replace(" ']", '').replace("['", '').replace("']", '')
    destination_data = 'DESLOCAMENTO'
    destination_data = destination_data + origin + '-' + destiny
    
    indices = [i for i, s in enumerate(tables_data) if 'PARA' in s]
    tables_data.insert(indices[0]+1, destination_data)

    return (tables_data)
            

fullText = getText(wordDoc)
tables_data = get_tables_data(wordDoc)

print(get_destination(tables_data))
